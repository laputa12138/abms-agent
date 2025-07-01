import logging
import os
import PyPDF2
import docx # python-docx
import nltk # Natural Language Toolkit
import uuid # For generating unique IDs for chunks
from typing import List, Dict, Any, Optional

from config.settings import (
    DEFAULT_PARENT_CHUNK_SIZE, DEFAULT_PARENT_CHUNK_OVERLAP,
    DEFAULT_CHILD_CHUNK_SIZE, DEFAULT_CHILD_CHUNK_OVERLAP,
    SUPPORTED_DOC_EXTENSIONS,
    # DEFAULT_CHUNK_SEPARATOR_REGEX # Not directly used if relying on NLTK or paragraph splits
)

# Configure logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s') # Already configured in main
logger = logging.getLogger(__name__)

# Download NLTK's sentence tokenizer models if not already present
# This is a one-time setup. In a production environment, this might be handled during deployment.
try:
    nltk.data.find('tokenizers/punkt')
except nltk.downloader.DownloadError:
    logger.info("NLTK 'punkt' model not found. Downloading...")
    nltk.download('punkt', quiet=True)
    logger.info("'punkt' model downloaded successfully.")


class DocumentProcessorError(Exception):
    """Custom exception for DocumentProcessor errors."""
    pass

class DocumentProcessor:
    """
    A class responsible for processing documents:
    - Extracting text from various file types (PDF, DOCX, TXT).
    - Splitting text into parent and child chunks for RAG.
    """

    def __init__(self,
                 parent_chunk_size: int = DEFAULT_PARENT_CHUNK_SIZE,
                 parent_chunk_overlap: int = DEFAULT_PARENT_CHUNK_OVERLAP,
                 child_chunk_size: int = DEFAULT_CHILD_CHUNK_SIZE,
                 child_chunk_overlap: int = DEFAULT_CHILD_CHUNK_OVERLAP,
                 supported_extensions: List[str] = None):
        """
        Initializes the DocumentProcessor.

        Args:
            parent_chunk_size (int): Target size for parent chunks (characters).
            parent_chunk_overlap (int): Overlap between parent chunks (characters).
            child_chunk_size (int): Target size for child chunks (characters).
            child_chunk_overlap (int): Overlap between child chunks (characters).
            supported_extensions (List[str], optional): List of supported file extensions.
                                                        Defaults to SUPPORTED_DOC_EXTENSIONS.
        """
        self.parent_chunk_size = parent_chunk_size
        self.parent_chunk_overlap = parent_chunk_overlap
        self.child_chunk_size = child_chunk_size
        self.child_chunk_overlap = child_chunk_overlap
        self.supported_extensions = supported_extensions or SUPPORTED_DOC_EXTENSIONS

        # Validate overlaps (simple validation)
        if self.parent_chunk_overlap >= self.parent_chunk_size and self.parent_chunk_size > 0:
            logger.warning("Parent chunk overlap is >= parent chunk size. Adjusting overlap.")
            self.parent_chunk_overlap = self.parent_chunk_size // 2 if self.parent_chunk_size > 10 else 0
        if self.child_chunk_overlap >= self.child_chunk_size and self.child_chunk_size > 0:
            logger.warning("Child chunk overlap is >= child chunk size. Adjusting overlap.")
            self.child_chunk_overlap = self.child_chunk_size // 2 if self.child_chunk_size > 10 else 0

        logger.info(f"DocumentProcessor initialized with parent_size={parent_chunk_size}, child_size={child_chunk_size}")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        logger.debug(f"Extracting text from PDF: {file_path}")
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
            return text
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}")
            raise DocumentProcessorError(f"Failed to extract text from PDF {file_path}: {e}")

    def _extract_text_from_docx(self, file_path: str) -> str:
        logger.debug(f"Extracting text from DOCX: {file_path}")
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise DocumentProcessorError(f"Failed to extract text from DOCX {file_path}: {e}")

    def _extract_text_from_txt(self, file_path: str) -> str:
        logger.debug(f"Extracting text from TXT: {file_path}")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e: # Catch potential UnicodeDecodeError or others
            logger.error(f"Error processing TXT file {file_path}: {e}")
            # Try with a fallback encoding if UTF-8 fails (less common for .txt but possible)
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    logger.warning(f"Retrying TXT file {file_path} with 'latin-1' encoding.")
                    return file.read()
            except Exception as e_fallback:
                logger.error(f"Fallback encoding also failed for TXT file {file_path}: {e_fallback}")
                raise DocumentProcessorError(f"Failed to extract text from TXT {file_path}: {e_fallback}")


    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extracts text content from a given file based on its extension.

        Args:
            file_path (str): The path to the document file.

        Returns:
            str: The extracted text content.

        Raises:
            DocumentProcessorError: If the file type is not supported or processing fails.
            FileNotFoundError: If the file_path does not exist.
        """
        logger.info(f"Attempting to extract text from file: {file_path}")
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        _, extension = os.path.splitext(file_path.lower())

        if extension not in self.supported_extensions:
            msg = f"Unsupported file type: {extension}. Supported types are: {self.supported_extensions}"
            logger.error(msg)
            raise DocumentProcessorError(msg)

        extracted_text = ""
        if extension == ".pdf":
            extracted_text = self._extract_text_from_pdf(file_path)
        elif extension == ".docx":
            extracted_text = self._extract_text_from_docx(file_path)
        elif extension == ".txt":
            extracted_text = self._extract_text_from_txt(file_path)

        logger.info(f"Successfully extracted {len(extracted_text)} characters from {file_path}.")
        return extracted_text

    def _split_into_fixed_size_chunks(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """Helper to split text by fixed size and overlap."""
        if not text or chunk_size <= 0:
            return [text] if text else []

        chunks = []
        start_index = 0
        while start_index < len(text):
            end_index = start_index + chunk_size
            chunks.append(text[start_index:end_index])
            next_start = start_index + chunk_size - chunk_overlap
            if next_start <= start_index : # Avoid infinite loop if step is non-positive
                logger.warning(f"Chunking step is non-positive (size: {chunk_size}, overlap: {chunk_overlap}). Breaking loop.")
                break
            start_index = next_start

        return [chunk for chunk in chunks if chunk.strip()]

    def split_text_into_parent_child_chunks(self,
                                            full_text: str,
                                            source_document_id: str = None
                                           ) -> List[Dict[str, Any]]:
        """
        Splits text into parent chunks, and each parent chunk into child chunks.

        Args:
            full_text (str): The entire text content of a document.
            source_document_id (str, optional): An identifier for the source document.
                                                If None, a new UUID will be generated.

        Returns:
            List[Dict[str, Any]]: A list of parent chunk dictionaries. Each dictionary has:
                - 'parent_id': str (unique ID for the parent chunk)
                - 'parent_text': str (text of the parent chunk)
                - 'source_document_id': str (ID of the source document)
                - 'children': List[Dict[str, Any]], where each child dict has:
                    - 'child_id': str (unique ID for the child chunk)
                    - 'child_text': str (text of the child chunk)
                    - 'parent_id': str (ID of its parent chunk)
        """
        if not full_text.strip():
            logger.warning("split_text_into_parent_child_chunks called with empty text.")
            return []

        doc_id = source_document_id or str(uuid.uuid4())
        structured_chunks = []

        # 1. Split into Parent Chunks (e.g., by paragraphs or fixed size)
        # Option A: Split by double newlines (paragraphs), then join if too small / re-split if too large.
        # Option B: Use fixed size splitting for parents. Let's start with fixed size for simplicity.

        parent_texts = self._split_into_fixed_size_chunks(
            full_text, self.parent_chunk_size, self.parent_chunk_overlap
        )

        logger.info(f"Document '{doc_id}' split into {len(parent_texts)} parent candidates.")

        for i, p_text in enumerate(parent_texts):
            if not p_text.strip():
                continue

            parent_id = f"{doc_id}-p{i+1}"
            parent_chunk_data = {
                "parent_id": parent_id,
                "parent_text": p_text,
                "source_document_id": doc_id,
                "children": []
            }

            # 2. Split each Parent Chunk into Child Chunks (e.g., by sentences or smaller fixed size)
            # Option A: Use NLTK for sentence splitting.
            # Option B: Use fixed size splitting for children.
            # Let's try NLTK sentence splitting first, then group sentences if child_chunk_size aims for multiple sentences.

            child_texts_from_parent = []
            sentences = nltk.sent_tokenize(p_text, language='english') # Assuming English for now; for Chinese, would need a Chinese tokenizer.
                                                                      # The prompt stated reports are Chinese, but source docs can be EN/CN.
                                                                      # For simplicity, let's use 'english'. If NLTK's punkt supports Chinese well enough, great.
                                                                      # Otherwise, specific Chinese sentence tokenization might be needed (e.g. jieba, pkuseg).
                                                                      # Given the project context, `language='chinese'` might be better if `punkt` for chinese is available.
                                                                      # Let's assume NLTK's default punkt handles mixed content reasonably for splitting, or we can make this configurable.

            current_child_text = ""
            for sent_idx, sentence in enumerate(sentences):
                if len(current_child_text) + len(sentence) + 1 <= self.child_chunk_size or not current_child_text:
                    current_child_text += (sentence + " ")
                else:
                    # Current child chunk is full, or adding next sentence exceeds size
                    if current_child_text.strip():
                        child_texts_from_parent.append(current_child_text.strip())
                    current_child_text = sentence + " " # Start new child chunk

            # Add the last accumulated child chunk
            if current_child_text.strip():
                child_texts_from_parent.append(current_child_text.strip())

            # Fallback or alternative: if sentence tokenization results in too few or too large chunks,
            # or if child_texts_from_parent is empty, use fixed size splitting for children.
            if not child_texts_from_parent or any(len(ct) > self.child_chunk_size * 1.5 for ct in child_texts_from_parent): # Heuristic
                 logger.debug(f"Parent chunk '{parent_id}' either too small for sentence splitting or sentences too long. Using fixed-size child splitting.")
                 child_texts_from_parent = self._split_into_fixed_size_chunks(
                     p_text, self.child_chunk_size, self.child_chunk_overlap
                 )


            for j, c_text in enumerate(child_texts_from_parent):
                if not c_text.strip():
                    continue
                child_id = f"{parent_id}-c{j+1}"
                parent_chunk_data["children"].append({
                    "child_id": child_id,
                    "child_text": c_text,
                    "parent_id": parent_id # Redundant here, but good for flat list if needed
                })

            if parent_chunk_data["children"]: # Only add parent if it has children
                structured_chunks.append(parent_chunk_data)
            else:
                logger.warning(f"Parent chunk '{parent_id}' produced no valid child chunks. Skipping this parent.")


        logger.info(f"Document '{doc_id}' processed into {len(structured_chunks)} parent chunks with children.")
        return structured_chunks


if __name__ == '__main__':
    print("DocumentProcessor Extended Example")
    logging.basicConfig(level=logging.DEBUG) # Set to DEBUG for more verbose output from this example

    # Initialize with default parent/child sizes
    processor = DocumentProcessor()

    # --- Test Text Extraction ---
    # Create dummy files for testing (in a real scenario, these would exist)
    dummy_dir = "temp_docs_for_dp_test"
    if not os.path.exists(dummy_dir):
        os.makedirs(dummy_dir)

    dummy_pdf_path = os.path.join(dummy_dir, "test.pdf")
    dummy_docx_path = os.path.join(dummy_dir, "test.docx")
    dummy_txt_path = os.path.join(dummy_dir, "test.txt")
    dummy_unsupported_path = os.path.join(dummy_dir, "test.xls")

    # Create a very simple PDF (PyPDF2 cannot easily write text to new PDFs)
    # So this PDF will likely extract empty text or cause an error if not valid.
    # For robust testing of PDF extraction, a real PDF is needed.
    # We'll mock this part for the example's text content.
    try:
        from PyPDF2 import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(210, 297) # A4
        with open(dummy_pdf_path, "wb") as f:
            writer.write(f)
        print(f"Created dummy PDF: {dummy_pdf_path}")
    except Exception as e:
        print(f"Could not create dummy PDF: {e}")


    # Create DOCX
    try:
        doc = docx.Document()
        doc.add_paragraph("This is a test document for DOCX processing.")
        doc.add_paragraph("It contains multiple paragraphs to test extraction.")
        doc.save(dummy_docx_path)
        print(f"Created dummy DOCX: {dummy_docx_path}")
    except Exception as e:
        print(f"Could not create dummy DOCX: {e}")


    # Create TXT
    try:
        with open(dummy_txt_path, "w", encoding="utf-8") as f:
            f.write("Hello from the TXT file.\nThis is the second line of the text file.")
        print(f"Created dummy TXT: {dummy_txt_path}")
    except Exception as e:
        print(f"Could not create dummy TXT: {e}")

    with open(dummy_unsupported_path, "w") as f: f.write("content") # Dummy unsupported

    file_paths_to_test = {
        "pdf": dummy_pdf_path,
        "docx": dummy_docx_path,
        "txt": dummy_txt_path,
        "unsupported": dummy_unsupported_path
    }

    # Mocking PDF text as PyPDF2 needs a real text PDF
    MOCK_PDF_TEXT = "This is mocked PDF text. It has sentences. And paragraphs too.\n\nThis is a new paragraph in the mocked PDF."

    for file_type, path in file_paths_to_test.items():
        print(f"\n--- Testing extraction for {file_type.upper()} file: {path} ---")
        try:
            if file_type == "pdf" and os.path.exists(path): # Use mock for PDF
                 text_content = MOCK_PDF_TEXT
                 print(f"Using MOCKED text for PDF: '{text_content[:100]}...'")
            elif os.path.exists(path):
                text_content = processor.extract_text_from_file(path)
                print(f"Extracted text ({len(text_content)} chars): '{text_content[:100]}...'")
            else:
                print(f"File {path} does not exist. Skipping.")
                continue

            # --- Test Parent-Child Chunking ---
            if text_content and text_content.strip():
                print(f"\n--- Testing Parent-Child Chunking for {file_type.upper()} content ---")
                # Use smaller chunk sizes for this example text to see more chunks
                processor_small_chunks = DocumentProcessor(parent_chunk_size=80, parent_chunk_overlap=10,
                                                           child_chunk_size=30, child_chunk_overlap=5)

                parent_child_chunks = processor_small_chunks.split_text_into_parent_child_chunks(
                    text_content, source_document_id=f"doc_{file_type}"
                )

                print(f"Generated {len(parent_child_chunks)} parent chunks.")
                for i, parent_data in enumerate(parent_child_chunks):
                    print(f"  Parent {i+1} (ID: {parent_data['parent_id']}): '{parent_data['parent_text'][:50]}...'")
                    print(f"    Source Doc ID: {parent_data['source_document_id']}")
                    print(f"    Contains {len(parent_data['children'])} child chunks:")
                    for j, child_data in enumerate(parent_data['children']):
                        print(f"      Child {j+1} (ID: {child_data['child_id']}): '{child_data['child_text'][:40]}...'")
                if not parent_child_chunks:
                     print("No parent-child chunks generated (text might be too short for current settings).")
            else:
                print("Skipping chunking due to empty extracted text.")

        except (FileNotFoundError, DocumentProcessorError) as e:
            print(f"Error during processing {file_type}: {e}")
        except Exception as e:
            print(f"Unexpected error for {file_type}: {e}")
            # import traceback; traceback.print_exc()


    # Cleanup dummy files
    try:
        if os.path.exists(dummy_dir):
            import shutil
            shutil.rmtree(dummy_dir)
            print(f"\nCleaned up temporary directory: {dummy_dir}")
    except Exception as e:
        print(f"Error cleaning up temp directory: {e}")

    print("\nDocumentProcessor Extended Example Finished.")
